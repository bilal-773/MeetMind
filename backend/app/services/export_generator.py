import os
import re
import io
from datetime import datetime
from fpdf import FPDF
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from bidi.algorithm import get_display

def has_urdu(text: str) -> bool:
    return bool(re.search(r"[\u0600-\u06ff]", text))

def reshape_text(text: str) -> str:
    reshaped_text = arabic_reshaper.reshape(text)
    return get_display(reshaped_text)

def strip_emojis(text: str) -> str:
    # Strip emojis and variation selectors
    emoji_pattern = re.compile(
        r'[\u2600-\u27BF]|\uFE0F|[\u2000-\u3300]|[\uD83C-\uDBFF][\uDC00-\uDFFF]',
        flags=re.UNICODE
    )
    return emoji_pattern.sub('', text)

def process_text(text: str) -> str:
    cleaned = strip_emojis(text)
    if has_urdu(cleaned):
        try:
            return reshape_text(cleaned)
        except Exception:
            return cleaned
    return cleaned

def format_duration(seconds: int) -> str:
    if not seconds:
        return "0s"
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    parts = []
    if h > 0:
        parts.append(f"{h}h")
    if m > 0:
        parts.append(f"{m}m")
    if s > 0 or not parts:
        parts.append(f"{s}s")
    return " ".join(parts)

def format_seconds_to_srt_time(seconds: float) -> str:
    hrs = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    ms = int(round((seconds - int(seconds)) * 1000))
    if ms >= 1000:
        ms = 999
    return f"{hrs:02d}:{mins:02d}:{secs:02d},{ms:03d}"

def get_system_font_paths():
    # Common Windows and Linux font locations
    candidates = [
        (r"C:\Windows\Fonts\tahoma.ttf", r"C:\Windows\Fonts\tahomabd.ttf"),
        (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    for reg, bold in candidates:
        if os.path.exists(reg):
            return reg, bold if os.path.exists(bold) else reg
    return None, None

class MeetingPDF(FPDF):
    def __init__(self, title_text, font_path=None):
        super().__init__()
        self.title_text = title_text
        self.font_path = font_path
        
    def header(self):
        self.set_y(10)
        if self.font_path:
            self.set_font("UnicodeFont", "B", 9)
        else:
            self.set_font("Helvetica", "B", 9)
        self.set_text_color(148, 163, 184) # text-muted
        self.cell(w=0, h=10, text="MeetMind AI - Meeting Intelligence", border=0, ln=0, align="L")
        self.cell(w=0, h=10, text=process_text(self.title_text[:50]), border=0, ln=1, align="R")
        self.line(10, 20, 200, 20)
        self.ln(5)
        
    def footer(self):
        self.set_y(-15)
        if self.font_path:
            self.set_font("UnicodeFont", "", 8)
        else:
            self.set_font("Helvetica", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(w=0, h=10, text="2026 all rights reserved. Made by Muhammad Bilal Asif", border=0, ln=0, align="L")
        self.cell(w=0, h=10, text=f"Page {self.page_no()}/{{nb}}", border=0, ln=1, align="R")

def set_pdf_font(pdf, style, size, font_path, bold_font_path):
    if font_path:
        pdf.set_font("UnicodeFont", style, size)
    else:
        pdf.set_font("Helvetica", style, size)

def write_markdown_text(pdf, text, font_path, bold_font_path, size=10, height=5.5):
    parts = text.split("**")
    is_bold = False
    for part in parts:
        clean_part = part.replace("*", "")
        if is_bold:
            set_pdf_font(pdf, "B", size, font_path, bold_font_path)
        else:
            set_pdf_font(pdf, "", size, font_path, bold_font_path)
        pdf.write(h=height, text=process_text(clean_part))
        is_bold = not is_bold

def write_markdown_pdf(pdf, text_content, font_path, bold_font_path):
    import re
    lines = text_content.split("\n")
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            pdf.ln(3)
            continue
            
        # Skip raw markdown tables (since we render a beautiful native table anyway)
        if line_strip.startswith("|"):
            continue
            
        # Horizontal Rule
        if line_strip == "---" or line_strip == "***":
            pdf.ln(2)
            pdf.set_draw_color(226, 232, 240)
            pdf.line(pdf.l_margin, pdf.y, 210 - pdf.r_margin, pdf.y)
            pdf.ln(3)
            continue

        # Headers
        if line_strip.startswith("# "):
            pdf.ln(4)
            set_pdf_font(pdf, "B", 16, font_path, bold_font_path)
            pdf.set_text_color(79, 70, 229) # Indigo
            txt = line_strip[2:]
            pdf.multi_cell(w=0, h=8, text=process_text(txt.replace("*", "")), ln=True)
            pdf.ln(2)
        elif line_strip.startswith("## "):
            pdf.ln(3)
            set_pdf_font(pdf, "B", 13, font_path, bold_font_path)
            pdf.set_text_color(14, 165, 233) # Sky Blue
            txt = line_strip[3:]
            pdf.multi_cell(w=0, h=7, text=process_text(txt.replace("*", "")), ln=True)
            pdf.ln(1)
        elif line_strip.startswith("### "):
            pdf.ln(2)
            set_pdf_font(pdf, "B", 11, font_path, bold_font_path)
            pdf.set_text_color(15, 23, 42) # Text primary
            txt = line_strip[4:]
            pdf.multi_cell(w=0, h=6, text=process_text(txt.replace("*", "")), ln=True)
            pdf.ln(1)
        # Bullet Points
        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            old_margin = pdf.l_margin
            pdf.set_left_margin(old_margin + 6)
            pdf.set_text_color(71, 85, 105) # Text secondary
            set_pdf_font(pdf, "", 10, font_path, bold_font_path)
            pdf.write(h=5.5, text="• ")
            write_markdown_text(pdf, line_strip[2:], font_path, bold_font_path, size=10)
            pdf.ln(5.5)
            pdf.set_left_margin(old_margin)
        # Numbered Lists
        elif re.match(r"^\d+\.\s+", line_strip):
            match = re.match(r"^(\d+\.\s+)(.*)$", line_strip)
            num_prefix = match.group(1)
            content = match.group(2)
            
            old_margin = pdf.l_margin
            pdf.set_left_margin(old_margin + 6)
            pdf.set_text_color(71, 85, 105)
            set_pdf_font(pdf, "B", 10, font_path, bold_font_path)
            pdf.write(h=5.5, text=num_prefix)
            write_markdown_text(pdf, content, font_path, bold_font_path, size=10)
            pdf.ln(5.5)
            pdf.set_left_margin(old_margin)
        # Quotes
        elif line_strip.startswith("> "):
            old_margin = pdf.l_margin
            pdf.set_left_margin(old_margin + 8)
            pdf.set_text_color(100, 116, 139) # text-muted
            write_markdown_text(pdf, line_strip[2:], font_path, bold_font_path, size=10)
            pdf.ln(5.5)
            pdf.set_left_margin(old_margin)
        else:
            pdf.set_text_color(71, 85, 105)
            write_markdown_text(pdf, line, font_path, bold_font_path, size=10)
            pdf.ln(5.5)

def generate_pdf(meeting: dict) -> bytes:
    """Generates a beautifully styled meeting minutes PDF using fpdf2."""
    title = meeting.get("title", "Meeting Minutes")
    font_path, bold_font_path = get_system_font_paths()
    
    pdf = MeetingPDF(title, font_path=font_path)
    pdf.alias_nb_pages()
    
    # Load Unicode Font if available
    if font_path:
        pdf.add_font("UnicodeFont", style="", fname=font_path)
        pdf.add_font("UnicodeFont", style="B", fname=bold_font_path)
        
    pdf.add_page()
    
    # Document Title Block
    set_pdf_font(pdf, "B", 20, font_path, bold_font_path)
    pdf.set_text_color(15, 23, 42) # Text primary
    pdf.multi_cell(w=0, h=10, text=process_text(title), ln=True, align="C")
    pdf.ln(6)
    
    # Metadata Box
    pdf.set_fill_color(248, 250, 252) # bg-secondary
    pdf.set_draw_color(226, 232, 240) # border
    set_pdf_font(pdf, "", 9, font_path, bold_font_path)
    pdf.set_text_color(71, 85, 105)
    
    # Duration, Speakers, Date
    duration_str = format_duration(meeting.get("duration_seconds", 0))
    speaker_count = meeting.get("speaker_count", 0)
    created_at = meeting.get("created_at")
    date_str = ""
    if isinstance(created_at, datetime):
        date_str = created_at.strftime("%B %d, %Y - %I:%M %p")
    elif isinstance(created_at, str):
        try:
            dt = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
            date_str = dt.strftime("%B %d, %Y - %I:%M %p")
        except Exception:
            date_str = created_at
            
    meta_text = (
        f"Date: {date_str}   |   "
        f"Duration: {duration_str}   |   "
        f"Speakers: {speaker_count}   |   "
        f"Languages: {', '.join(meeting.get('languages_detected', ['EN'])).upper()}"
    )
    # Draw border box around metadata
    pdf.cell(w=0, h=8, text=process_text(meta_text), border=1, ln=1, align="C", fill=True)
    pdf.ln(8)
    
    # Section: Meeting Minutes
    minutes_content = meeting.get("minutes_en") or meeting.get("minutes") or "No minutes generated."
    write_markdown_pdf(pdf, minutes_content, font_path, bold_font_path)
    pdf.ln(5)
    
    # Section: Action Items
    action_items = meeting.get("action_items", [])
    if action_items:
        pdf.ln(4)
        set_pdf_font(pdf, "B", 13, font_path, bold_font_path)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(w=0, h=7, text="Action Items Table", border=0, ln=1)
        pdf.ln(2)
        
        # Table Headers
        pdf.set_fill_color(238, 242, 255) # accent-light
        set_pdf_font(pdf, "B", 9, font_path, bold_font_path)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(w=90, h=7, text="Task", border=1, ln=0, fill=True)
        pdf.cell(w=35, h=7, text="Owner", border=1, ln=0, fill=True)
        pdf.cell(w=35, h=7, text="Deadline", border=1, ln=0, fill=True)
        pdf.cell(w=25, h=7, text="Priority", border=1, ln=1, fill=True)
        
        # Table Rows
        set_pdf_font(pdf, "", 9, font_path, bold_font_path)
        pdf.set_text_color(71, 85, 105)
        for item in action_items:
            # Handle model object vs dictionary mapping
            task = item.task if hasattr(item, "task") else item.get("task", "")
            owner = item.owner if hasattr(item, "owner") else item.get("owner", "")
            deadline = item.deadline if hasattr(item, "deadline") else item.get("deadline", "")
            priority = item.priority if hasattr(item, "priority") else item.get("priority", "medium")
            
            task = task or "-"
            owner = owner or "-"
            deadline = deadline or "-"
            priority = priority or "medium"
            
            # Simple column cell wraps
            # FPDF2 supports multi_cell or simple clipping
            pdf.cell(w=90, h=7, text=process_text(task[:45]), border=1, ln=0)
            pdf.cell(w=35, h=7, text=process_text(str(owner)[:18]), border=1, ln=0)
            pdf.cell(w=35, h=7, text=process_text(str(deadline)[:18]), border=1, ln=0)
            pdf.cell(w=25, h=7, text=process_text(str(priority).upper()), border=1, ln=1)
            
    # Return document bytes
    out_buf = io.BytesIO()
    pdf.output(out_buf)
    return out_buf.getvalue()

def add_markdown_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    parts = text.split("**")
    is_bold = False
    for part in parts:
        run = p.add_run(part)
        if is_bold:
            run.bold = True
        is_bold = not is_bold
    return p

def generate_docx(meeting: dict) -> bytes:
    """Generates a professionally formatted Word DOCX document."""
    title = meeting.get("title", "Meeting Minutes")
    doc = docx.Document()
    
    # Document Style Settings
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    # Metadata Table
    duration_str = format_duration(meeting.get("duration_seconds", 0))
    speaker_count = meeting.get("speaker_count", 0)
    created_at = meeting.get("created_at")
    date_str = ""
    if isinstance(created_at, datetime):
        date_str = created_at.strftime("%B %d, %Y - %I:%M %p")
    elif isinstance(created_at, str):
        try:
            dt = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
            date_str = dt.strftime("%B %d, %Y - %I:%M %p")
        except Exception:
            date_str = created_at
            
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Duration'
    hdr_cells[2].text = 'Speakers'
    hdr_cells[3].text = 'Languages'
    
    row_cells = table.add_row().cells
    row_cells[0].text = date_str
    row_cells[1].text = duration_str
    row_cells[2].text = f"{speaker_count} identified"
    row_cells[3].text = ", ".join(meeting.get("languages_detected", ["EN"])).upper()
    
    doc.add_paragraph() # Spacing
    
    # Section: Meeting Minutes
    doc.add_heading("Meeting Minutes", level=1)
    minutes_content = meeting.get("minutes_en") or meeting.get("minutes") or "No minutes generated."
    
    # Map markdown content to Word blocks
    for line in minutes_content.split("\n"):
        line_strip = line.strip()
        if not line_strip:
            continue
        if line_strip.startswith("# "):
            doc.add_heading(line_strip[2:], level=1)
        elif line_strip.startswith("## "):
            doc.add_heading(line_strip[3:], level=2)
        elif line_strip.startswith("### "):
            doc.add_heading(line_strip[4:], level=3)
        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            add_markdown_paragraph(doc, line_strip[2:], style='List Bullet')
        elif line_strip.startswith("> "):
            p = add_markdown_paragraph(doc, line_strip[2:])
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 116, 139)
        else:
            add_markdown_paragraph(doc, line)
            
    # Section: Action Items
    action_items = meeting.get("action_items", [])
    if action_items:
        doc.add_paragraph()
        doc.add_heading("Action Items", level=1)
        
        act_table = doc.add_table(rows=1, cols=4)
        act_table.style = 'Medium Shading 1 Accent 1'
        act_hdr = act_table.rows[0].cells
        act_hdr[0].text = 'Task'
        act_hdr[1].text = 'Owner'
        act_hdr[2].text = 'Deadline'
        act_hdr[3].text = 'Priority'
        
        for item in action_items:
            task = item.task if hasattr(item, "task") else item.get("task", "")
            owner = item.owner if hasattr(item, "owner") else item.get("owner", "")
            deadline = item.deadline if hasattr(item, "deadline") else item.get("deadline", "")
            priority = item.priority if hasattr(item, "priority") else item.get("priority", "medium")
            
            r_cells = act_table.add_row().cells
            r_cells[0].text = task or "-"
            r_cells[1].text = str(owner or "-")
            r_cells[2].text = str(deadline or "-")
            r_cells[3].text = str(priority or "medium").upper()
            
    # Section: Transcript Log
    transcript = meeting.get("transcript") or meeting.get("transcript_raw")
    if transcript:
        doc.add_paragraph()
        doc.add_heading("Transcript Log", level=1)
        
        # Build speaker display map
        speaker_map = {}
        for spk in meeting.get("speakers", []):
            spk_key = spk.speaker_key if hasattr(spk, "speaker_key") else spk.get("speaker_key")
            disp_name = spk.display_name if hasattr(spk, "display_name") else spk.get("display_name")
            if spk_key and disp_name:
                speaker_map[spk_key] = disp_name
                
        for seg in transcript[:250]: # Limit to first 250 dialogue rows for readability
            start_s = seg.get("start") if isinstance(seg, dict) else getattr(seg, "start", 0)
            speaker_key = seg.get("speaker") if isinstance(seg, dict) else getattr(seg, "speaker", "")
            text = seg.get("text") if isinstance(seg, dict) else getattr(seg, "text", "")
            
            speaker_name = speaker_map.get(speaker_key, speaker_key or "Unknown")
            time_str = format_seconds_to_srt_time(start_s).split(',')[0]
            
            p = doc.add_paragraph()
            r_time = p.add_run(f"[{time_str}] ")
            r_time.font.color.rgb = RGBColor(14, 165, 233)
            r_time.font.name = 'Courier New'
            
            r_spk = p.add_run(f"{speaker_name}: ")
            r_spk.font.bold = True
            r_spk.font.color.rgb = RGBColor(79, 70, 229)
            
            p.add_run(text)
            
    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue()

def generate_srt(transcript: list, speakers: list = None) -> bytes:
    """Generates standard SRT Subtitle bytes from transcript segments."""
    if not transcript:
        return b""
        
    # Build speaker map
    speaker_map = {}
    if speakers:
        for spk in speakers:
            spk_key = spk.speaker_key if hasattr(spk, "speaker_key") else spk.get("speaker_key")
            disp_name = spk.displayName if hasattr(spk, "displayName") else spk.get("display_name")
            if spk_key and disp_name:
                speaker_map[spk_key] = disp_name
                
    srt_lines = []
    for idx, seg in enumerate(transcript):
        start_s = seg.get("start") if isinstance(seg, dict) else getattr(seg, "start", 0)
        end_s = seg.get("end") if isinstance(seg, dict) else getattr(seg, "end", 0)
        speaker_key = seg.get("speaker") if isinstance(seg, dict) else getattr(seg, "speaker", "")
        text = seg.get("text") if isinstance(seg, dict) else getattr(seg, "text", "")
        
        speaker_name = speaker_map.get(speaker_key, speaker_key or "Unknown")
        time_start = format_seconds_to_srt_time(start_s)
        time_end = format_seconds_to_srt_time(end_s)
        
        # Subtitle Block
        srt_lines.append(str(idx + 1))
        srt_lines.append(f"{time_start} --> {time_end}")
        srt_lines.append(f"[{speaker_name}]: {text}")
        srt_lines.append("") # Blank separator line
        
    srt_content = "\n".join(srt_lines)
    return srt_content.encode("utf-8")
